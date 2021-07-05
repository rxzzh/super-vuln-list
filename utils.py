from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.shared import Cm
import re

def gadget_fill_cell(row, fields: list):
    for i in range(len(fields)):
        # set row height
        row.height = Cm(0.8)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        # set vertical center
        row.cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # set horizontal center
        paragraph = row.cells[i].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # write content to cell
        paragraph.add_run(str(fields[i]))

def doc_add_comment(doc, comment: str):
    doc.add_paragraph().text = comment

ip_regex = re.compile('^([0-9]+\.){3}[0-9]+.html$')
true_ip_regex = re.compile('^([0-9]+\.){3}[0-9]+$')

def concat_path(path_head, path_tail):
    path_head = path_head.strip('/')
    path_tail = path_tail.lstrip('/')
    return path_head + '/' + path_tail

def gadget_fill_cell_super(cells, fields):
    for i in range(len(fields)):
        cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph = cells[i].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run(str(fields[i]))

def gadget_set_row_height(rows):
    for row in rows:
        row.height = Cm(0.8)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST